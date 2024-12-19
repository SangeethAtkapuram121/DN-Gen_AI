# import fitz  # PyMuPDF

# def remove_header_footer(input_pdf, output_pdf, header_height, footer_height):
#     # Open the PDF
#     pdf_document = fitz.open(input_pdf)
#     num_pages = pdf_document.page_count

#     # Create a new PDF for the output
#     output_document = fitz.open()

#     for page_number in range(num_pages):
#         # Get the current page
#         page = pdf_document.load_page(page_number)
#         rect = page.rect
#         new_rect = fitz.Rect(rect.x0, rect.y0 + header_height, rect.x1, rect.y1 - footer_height)
        
#         # Create a new page with the adjusted rectangle
#         new_page = output_document.new_page(width=new_rect.width, height=new_rect.height)
#         new_page.show_pdf_page(new_rect, pdf_document, page_number, clip=new_rect)

#     # Save the new PDF
#     output_document.save(output_pdf)
# # Usage
# input_pdf = "Department of CSE-AIML Faculty.pdf"
# output_pdf = "./hnf_removed/Department of CSE-AIML Faculty.pdf"
# header_height = 60  # Adjust based on the height of the header
# footer_height = 10  # Adjust based on the height of the footer
# remove_header_footer(input_pdf, output_pdf, header_height, footer_height)
# pdf_file_path = "Department of CSE-AIML Faculty.pdf"
doc_path = './cleaned_docx/Department of CSE-AIML Faculty.docx'

# from pdf2docx import Converter

# def convert_pdf_to_docx(pdf_file_path,doc_path,start_page_index):   
    
#     # Using the built-in function, convert the PDF file to a document file by saving it in a variable.
#     cv = Converter(pdf_file_path)
    
#     # Storing the Document in the variable's initialised path
#     cv.convert(doc_path, start=start_page_index)
    
#     # Conversion closure through the function close()
#     cv.close()
    
# convert_pdf_to_docx(pdf_file_path,doc_path,38)

from docx import Document
from docx.shared import Pt
import re

def count_dots(first_word):
    dot_count = 0
    for character in first_word:
        if character == ".":
            dot_count += 1 
    
    return dot_count


def add_space_between_numbers_and_alphabets(input_string):
    # Using regex to find positions where a number is followed by an alphabet or vice versa
    result = re.sub(r'(\d)([a-zA-Z])', r'\1 \2', input_string)
    result = re.sub(r'([a-zA-Z])(\d)', r'\1 \2', result)
    return result

def set_heading_by_font_size(doc_path):
    # Load the document
    doc = Document(doc_path)
    
    # Iterate over all paragraphs
    for para in doc.paragraphs:
        # Check if the paragraph has runs with specific font sizes
        if para.runs:
            for run in para.runs:
                if run.font.size:
                    font_size_pt = run.font.size.pt
                    if font_size_pt >= 18:
                        para.style = 'Heading 1'
                        break  # Once the style is set, break out of the loop
                    elif font_size_pt >= 14:
                        first_word = run.text.split()[0]
                        first_word = add_space_between_numbers_and_alphabets(first_word)
                        if first_word[-1] == ".":
                            first_word = first_word[:-1]
                        dots = count_dots(first_word)
                        if dots == 1:
                            para.style = 'Heading 2'
                            break  # Once the style is set, break out of the loop
                        elif dots == 2:
                            para.style = 'Heading 3'
                            break
                        # elif dots == 3 or dots == 0 :
                        #     para.style = 'Heading 4'
                        #     break
                        # else:
                        #     para.style = 'Heading 5'
                        #     break
    # Save the document
    doc.save(doc_path)    

# Example usage
set_heading_by_font_size("./cleaned_docx/Department of CSE-AIML Faculty.docx")


from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import xml.etree.ElementTree as ET
import time
import os
import re
from lxml import etree
# Import the required modules
from pdf2docx import Converter
from docx.enum.table import WD_ALIGN_VERTICAL
 
from docx.shared import Pt
from docx.oxml.text.paragraph import CT_P
from pathlib import Path
import markdownify

### Converts a docx file to html file. 
def convert_docx_to_html(doc_path,html_file,footer_string):
    # dir_path = "./SOURCEDIR"
    # file = "SAMPLE.docx"
    # image_path will be a directory to store image files; if the directory does not exist it will be created

    document = Document(doc_path)
    # for section in document.sections:
    #     footer = section.footer
    #     footer.paragraphs[0].text  = footer.paragraphs[0].text.replace("Root Cause Analysis Confidential", "")

    body = "<html><body>"
    list_items = []
 
    ### iterates over an document elements to generate HTML content. Handles lists, headings and paragraphs ensuring proper HTML tags are used.
    ### writes the resulting html content to a specified file
    for block in iter_block_items(document):
        if isinstance(block, Paragraph) and block.text.replace("\n", "").strip() != footer_string :
            tmp_heading_type = get_heading_type(block)
            if re.match("List", tmp_heading_type):
                list_items.append("<li>" + block.text + "</li>")
            elif re.match("Paragraph", tmp_heading_type):
                list_items.append("<p>" + block.text + "</p>")
            
            else:
                #images = render_image(document, block, image_path, image_path)
                if len(list_items) > 0:
                    body += render_list_items(list_items)
                    list_items = []
                #if len(images) > 0:
                #    body = body + images
                else:
                    # modified to use a different outer_tag if a 'Heading' style is found in the original paragraph
                    if 'Heading' in tmp_heading_type:
                        outer_tag = 'h' + tmp_heading_type.split(' ')[-1]
                    else:
                        outer_tag = 'p'
                    body = body + render_runs(block.runs, outer_tag)
        elif isinstance(block, Table):
            body = render_table(block, document, body)
   
    body += "</body></html>"
    # body = normalize('NFKD', body).encode('ascii','ignore')
    with open(html_file,"wb") as file:
        file.write(body.encode('utf8', 'ignore'))
   
    # if os.path. exists(doc_path):
    #     os. remove(doc_path)
 
    return body
 
### Yields paragraphs and tables in a given parent element(either a document or a table cell)
### Facilitates iterating over different block items for processing  
def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)
 
### prints the content of a table to the console for debugging purposes        
def table_print(block):
    table = block
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                print(paragraph.text,'  ',end='')
        print("\n")
 
### checks if the first row or the last row of a table has a top or bottom border
### Determines table formatting based on border presence
def has_top_border(table):
    # if len(table.rows) >= 2:
    first_row = table.rows[0]._element
    check_if_borders_exist=first_row.find('.//w:tcPr/w:tcBorders', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
   
    if (check_if_borders_exist is not None):
        top_border_exists = first_row.find('.//w:tcPr/w:tcBorders/w:top', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    else:
        return True
   
    if(top_border_exists is None):
        return None
    else:
        return top_border_exists is not None
    # return False
 
def has_bottom_border(table):
    # if len(table.rows) >= 2:
    last_row = table.rows[-1]._element
   
    check_if_borders_exist=last_row.find('.//w:tcPr/w:tcBorders', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    if (check_if_borders_exist is not None):
        bottom_border_exists = last_row.find('.//w:tcPr/w:tcBorders/w:bottom', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    else:
        return True
    return bottom_border_exists is not None
    # return False  
 
### Finds and returns the paragraph preceding a specified table that contains a particular string
def find_paragraph_before_table(doc, table,search_str):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.text.strip() == search_str:#table.cell(0, 0).text.strip():
                return paragraph
    return None
 
### Inserts a new paragraph after a specified index in the documents paragraphs.
def insert_paragraph_after(doc,paragraphs, idx, text=None):
    next_paragraph_idx = idx + 1
    if idx == len(paragraphs):
        return doc.add_paragraph(text)
    next_paragraph = paragraphs[next_paragraph_idx]
    return next_paragraph.insert_paragraph_before(text)
 
### Adds border to all cells in a table.
def add_border_to_table(table):
    # Add content to the table
    for row in table.rows:
        for cell in row.cells:
            # Set text and alignment
            cell.text = ""
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
 
            # Set border properties
            for border in cell._element.xpath('.//w:tcBorders/w:bottom'):
                border.attrib.clear()
                border.attrib['w:val'] = 'single'
                border.attrib['w:sz'] = '6'
                border.attrib['w:color'] = '000000'
    return table
 
### moves a table to immediately follow a specified paragraph                
def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)
 
def is_cell_text_bold(cell):
    if cell.paragraphs:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if run.bold:
                    return True
    return False
 
### splits a table into two tables at a specified row index. Inserts a new paragraph between the split tables
def split_table(doc, table_index, split_row_index,split_table_title,paragraph_string):
    original_table = doc.tables[table_index]
 
    # Create two new tables
    table1 = doc.add_table(rows=0, cols=len(original_table.columns))
    table2 = doc.add_table(rows=0 , cols=len(original_table.columns))  
 
    table1.style = 'Table Grid'
    table2.style = 'Table Grid'
    # Copy data to the new tables
    for i, row in enumerate(original_table.rows):
        new_table = table1 if i <= split_row_index else table2
        new_row = new_table.add_row().cells
        for j, cell in enumerate(row.cells):
            if (j<len(row._tr.tc_lst)):
                new_row[j].text = cell.text
                if is_cell_text_bold(cell):
                    paragraph = new_row[j].paragraphs[0]
                    run = paragraph.runs[0]
                    run.bold=True
 
    # Find the paragraph preceding the original table
    original_paragraph = find_paragraph_before_table(doc, original_table,split_table_title)
 
    # Insert the new tables before the paragraph containing the original table
    if original_paragraph:
        paragraphs = list(doc.paragraphs)
        for idx, paragraph in enumerate(paragraphs):
            if paragraph.text.strip() == split_table_title:
                insert_paragraph_after(doc,paragraphs, idx, paragraph_string)
                break
       
        new_original_paragraph = find_paragraph_before_table(doc, original_table,paragraph_string)
        move_table_after(table2, new_original_paragraph)  
        move_table_after(table1, original_paragraph)
 
        # original_paragraph.insert_paragraph_after(table1._element)
        # original_paragraph.insert_paragraph_after(table2._element)
 
    # Remove the original table from the document
    original_table._element.getparent().remove(original_table._element)
 
### Generates html for a table , including handling nested tables and images within cells
# Modified to treat cell content as a set of blocks to process
def render_table(block,document,body):
    table = block
    html=""
    borders = has_top_border(table)
    if borders is not None and borders==True:
        html = "<table class='table table-bordered' border='1' cellspacing='0' >"
    else:
        while re.search(r'<p>\s*</p>$', body):
            body = re.sub(r'<p>\s*</p>$', '', body)
        if body.endswith("</table>"):
            body = body[:-len("</table>")]
   
    for row in table.rows:
        if borders is not None:
            html += "<tr>"
        for cell in row.cells:
             if borders is not None:
                html += "<td>"
             cbody = ""
             clist_items = []
             for cblock in iter_block_items(cell):
                if isinstance(cblock, Paragraph):
                    tmp_heading_type = get_heading_type(cblock)
                    if re.match(r"List\sParagraph",tmp_heading_type):
                        clist_items.append("<li>" + cblock.text + "</li>")
                    else:
                        #images = render_image(document,cblock,image_path,image_path)
                        if len(clist_items) > 0:
                            cbody += render_list_items(clist_items)
                            clist_items = []
                        #if len(images) > 0:
                            #cbody = cbody + images
                        else:
                            cbody = cbody + render_runs(cblock.runs)
                elif isinstance(cblock, Table):
                    cbody += render_table(cblock,document,"")
             html += cbody + " "
             if borders is not None:
                html += "</td>"
       
        if borders is not None:
            html += "</tr>"
       
    if (has_bottom_border(table)==True):        
         html += "</table>"
   
    body+=html
    return body
   
# Modified to use a different outer_tag if a 'Heading' style is found in the original paragraph
def render_runs(runs, outer_tag='p'):
    # Initialize the HTML with the outer tag
    html = "<" + outer_tag + ""
 
    # Check for bold text in any run
    for run in runs:
        if run.bold:
            html+=" style='font-weight:bold'"
            break
 
    html+=">"
 
    # Check if the runs belong to a list
    is_list = False
    list_items = []
   
    for run in runs:
        text = run.text.strip()
        if re.match(r'^(\d+\.\s+|\*\s+|-\s+)', text):
            is_list = True
            list_items.append("<li>" + text + "</li>")
        else:
            html += text
 
    if is_list:
        html = "<ul>" + "".join(list_items) + "</ul>"
    else:
        html += "</" + outer_tag + ">  "
   
    return html
 
### Converts a list of items into an HTML unordered list(<ul>)  
def render_list_items(items):
    html = "<ul>"
    for item in items:
        html += item
    html += "</ul>"
    return html
   
 ### returns the style name of a paragraph block block,used to determine heading levels
def get_heading_type(block):
    return block.style.name
 
 
### Converts to a PDF file to a DOCX file starting from a specified page index using pdf2docx li    
def convert_pdf_to_docx(pdf_file_path,doc_path,start_page_index):  
   
    # Using the built-in function, convert the PDF file to a document file by saving it in a variable.
    cv = Converter(pdf_file_path)
   
    # Storing the Document in the variable's initialised path
    cv.convert(doc_path,start=start_page_index)
   
    # Conversion closure through the function close()
    cv.close()


def extract_text_from_nested_table(cell):
    """
    Extracts all the text from a cell, including text from any nested tables.

    creates an empty list to append the extracted text from the table cell data(the notes , block sections inside the <td> cells)

    if found a paragraph appends it to the cell as it is.
    
    """
    extracted_text = []
    for paragraph in cell.paragraphs:
        extracted_text.append(paragraph.text)
   
    for table in cell.tables:
        for row in table.rows:
            row_text = []
            for nested_cell in row.cells:
                nested_text = extract_text_from_nested_table(nested_cell)
                row_text.append(nested_text)
            extracted_text.append(' | '.join(row_text))  # Join columns with a separator
   
    return ' '.join(extracted_text).strip()

# Removing Stop Words


import nltk
import shutil
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from docx import Document
nltk.download('stopwords')
nltk.download('punkt')
 
def remove_stop_words(text):

    ''' 
    Removes stop words using nltk library and tokenization
    '''
    # Load stop words
    stop_words = stopwords.words('english')
    stop_words.append('VynamicView-ProView')
    stop_words = set(stop_words)
    stop_words.remove("no")
    stop_words.remove("not")
    # Tokenize the text
    tokens = word_tokenize(text)
    tokens = list(map(lambda x: add_space_between_numbers_and_alphabets(x), tokens))
    # Filter out the stop words
    filtered_tokens = [w for w in tokens if not w in stop_words]
    # Join the tokens back into a string
    filtered_text = ' '.join(filtered_tokens)
    return filtered_text

def process_docx(docx_path):
    # Load the .docx file
    doc = Document(docx_path)
    # Remove stop words from paragraphs
    for para in doc.paragraphs:
        para.text = remove_stop_words(para.text)
    # Remove stop words from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = remove_stop_words(cell.text)
    
    # Save the modified document

    ''' 
    Changes are saved to the same file path
    '''
    stp_wrd = docx_path
    doc.save(stp_wrd)

    print("Removed stop words")


def stop(filename):
    # Path to the .docx file
    docx_path = filename
 
    # Process the document
    process_docx(docx_path)
    



# Also extracts the text in the blocks which are inside the table cells
def flatten_table(table):
    """
    Flattens a table by extracting text from nested tables and returning rows of text.
    """
    flattened_rows = []
    for row in table.rows:
        flattened_row = []
        for cell in row.cells:
            cell_text = extract_text_from_nested_table(cell)
            flattened_row.append(cell_text.strip())
        flattened_rows.append(flattened_row)
    return flattened_rows
def remove_nested_tables_from_table(table):
    """
    Removes nested tables from a table by replacing them with flattened text.
    """
    flattened_rows = flatten_table(table)
    new_table = Document().add_table(rows=len(flattened_rows), cols=len(flattened_rows[0]))
    for row_idx, row in enumerate(flattened_rows):
        for col_idx, cell_text in enumerate(row):
            new_table.cell(row_idx, col_idx).text = cell_text
    return new_table
 
def nested_docx(input_path):
    """
    Processes a .docx file, simplifying nested tables and saving the output.
    """
    doc = Document(input_path)
    new_doc = Document()
 
    for element in doc.element.body:
        if element.tag.endswith('tbl'):
            # Find the table in the original document
            for table in doc.tables:
                if table._element == element:
                    new_tbl = remove_nested_tables_from_table(table)
                    new_doc.element.body.append(new_tbl._element)
                    break
        else:
            new_doc.element.body.append(element)
    
    ''' 
    Changes are saved to the same file path
    '''
    new_doc.save(input_path)
    print("Simplifies nested table")
    
    #stop(input_path)

 



# def related_info(filename):
 
#     # Load the document
#     doc = Document(filename)

#     ''' 
#     Removes the Related Information in the .docx files
#     logic used: Remove information until the next heading(Any heading) is seen.
    
#     '''
 
#     # Initialize a flag to mark the start of the 'Related information' section
#     in_related_info_section = False
#     # Initialize a variable to store the paragraph to be removed
#     heading_to_remove = None
 
#     # Iterate through the paragraphs to find and mark the 'Related information' section
#     for para in doc.paragraphs:
#         if 'Related information' in para.text or 'Related Information' in para.text or 'Related documentation' in para.text:
#             # Mark the start of the section and store the heading paragraph
#             in_related_info_section = True
#             heading_to_remove = para
#         elif para.style.name.startswith('Heading') and in_related_info_section:
#             # If another heading is found, stop removing content
#             in_related_info_section = False
#         elif in_related_info_section:
#             # Add the paragraph to the list of paragraphs to be removed
#             if heading_to_remove:
#                 # Remove the stored heading paragraph
#                 p = heading_to_remove._element
#                 p.getparent().remove(p)
#                 p._p = p._element = None
#                 heading_to_remove = None
#             # Remove the current paragraph
#             p = para._element
#             p.getparent().remove(p)
#             p._p = p._element = None
 
#     # Save the modified document.

#     ''' 
#     Changes are saved to the same file path
#     '''
#     doc.save(filename)
#     print("Removed related information")
    
#     #nested_docx(filename)


from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
 
def convert_single_cell_tables_to_text(docx_path):
    # Load the document
    doc = Document(docx_path)

    ''' 
    Processes the blocks into text 
    logic used: Considered blocks as a table of single row with single data cell.
    
    '''
 
    # Collect tables to process in a separate list to avoid modifying the document structure during iteration
    tables_to_replace = []
 
    # Find all tables and check if they are single row and single cell
    for table in doc.tables:
        if len(table.rows) == 1 and len(table.columns) == 1:
            cell_text = table.cell(0, 0).text
            tables_to_replace.append((table, cell_text))
 
    # Replace each identified table with its text
    for table, cell_text in tables_to_replace:
        # Create a new paragraph with the extracted text
        new_paragraph = OxmlElement('w:p')
        new_run = OxmlElement('w:r')
        new_text = OxmlElement('w:t')
        new_text.text = cell_text
 
        new_run.append(new_text)
        new_paragraph.append(new_run)
 
        # Replace the table with the new paragraph
        table._element.getparent().replace(table._element, new_paragraph)

    print(docx_path)
    # Save the modified document

    ''' 
    Changes are made on the same file path
    '''
    doc.save(docx_path)
    print("Converted all the blocks into text")
    
    #related_info(new_docx_path)
 


import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

 
def remove_images_from_docx(input_path, images_folder):
    # Load the document
    
    ''' 
    Removes images from .docx files.

    Saves these images in a folder which is named same as the .docx file name.
    
    '''

    doc = Document(input_path)

    ''' 
    Uses regular expressions to split the input path by character "/" .

    A list called splitname is formed which will have n number of list items.

    Anyway second item of list is selected as folder name which is same as .docx file name
    
    '''
    splitname = re.split('[/ ]', input_path)


    if len(splitname) > 2:
        image_path = ''.join(splitname[1:2]) + '-images'

    
    else:
        image_path = splitname[1][:-5] + '-images'

    # Create the images folder if it doesn't exist
    # dest='/'+''

    final_folder=images_folder+'/'+image_path


    if not os.path.exists(final_folder):
        os.makedirs(final_folder)
    
    def save_image(img_part, img_counter):

        ''' 
        Saves the image in 
        
        '''
        img_data = img_part.blob
        img_format = img_part.content_type.split('/')[-1]
        img_name = f'image_{img_counter}.{img_format}'
        img_path = os.path.join(final_folder, img_name)
        with open(img_path, 'wb') as img_file:
            img_file.write(img_data)
        return img_path
    
    img_counter = 1
 
    # Remove image elements from paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            drawing_elements = run._element.xpath('.//w:drawing')
            for drawing in drawing_elements:
                blip = drawing.xpath('.//a:blip')
                if blip:
                    rId = blip[0].get(qn('r:embed'))
                    if rId in doc.part.rels:
                        rel = doc.part.rels[rId]
                        if "image" in rel.reltype:
                            save_image(rel.target_part, img_counter)
                            img_counter += 1
                            drawing.getparent().remove(drawing)
 
    # Remove image elements from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        drawing_elements = run._element.xpath('.//w:drawing')
                        for drawing in drawing_elements:
                            blip = drawing.xpath('.//a:blip')
                            if blip:
                                rId = blip[0].get(qn('r:embed'))
                                if rId in doc.part.rels:
                                    rel = doc.part.rels[rId]
                                    if "image" in rel.reltype:
                                        save_image(rel.target_part, img_counter)
                                        img_counter += 1
                                        drawing.getparent().remove(drawing)
 
    # Save the modified document

    file_name = os.path.basename(input_path)

    dest="./cleaned_docx/"
    output_docx_path=dest+file_name

    ''' 
    changes saved to different document in cleaned_docx folder
    '''
    doc.save(output_docx_path)
    
    print("image retrieval is done")
    
    ''' 
    function call is made here itself because returning a string to the __main__ gives NoneType Value as result
    '''

    convert_single_cell_tables_to_text(output_docx_path)

import os
import shutil
def iterate_retrieve_images():

    ''' 
    Iterates over files in docx folder
    
    The path of the file and the destination folder for images retrieval 
    of .docx file are input arguments to function remove_images_from_docx()  
    
    '''

    mypath1 = './cleaned_docx/'
    l = [f for f in os.listdir(mypath1) if f.endswith('.docx')]
    for i in l:
        print("Document "+i+"is now iterating in the data preprocessing pipeline")
        remove_images_from_docx(mypath1+i,'./cleaned_docx/retrieved_images')

import pathlib
def extract_html_md():

    '''
    
    Define cleaned documents folder path for src_doc_path
    Define the directory in which .html resides for dest_path
    Define the directory in which .md files reside for md_dest_path

    '''

    src_doc_path='./cleaned_docx/'
    dest_path='./html_processed/'
    md_dest_path='./md_processed/'

    dirs = os.listdir( src_doc_path )
    docs=[]
    for file in dirs:

        ''' 

        Checks if the file is a .docx file and specifies file names for .html and .md files.
        Converts .docx files into .md files using .html fileswhich is required for further processes like Generating Index

        '''
        if file.endswith('.docx'):
            html_file = dest_path+Path(file).stem+".html"
            md_file=md_dest_path+Path(file).stem+".md"
            html_text= convert_docx_to_html(src_doc_path+file, html_file,"")
            markdown_text = markdownify.markdownify(html_text, heading_style="ATX")
            with open(md_file,"wb") as file:
                file.write(markdown_text.encode('utf8', 'ignore'))


if __name__ == "__main__":

    ''' 
    Excecutes all the required functions for Data Preprocessing.
    The following is the workflow

    1. Retrieves images from .docx files
    2. Stores the resulted .docx file in cleaned_docx folder
    3. Converts blocks into text
    4. Removes Related information
    5. Simplifies nested tables including the block content in the table cells
    6. Removes stop words
    7. Converts .docx files into .html files then into .md files
    
    '''

    iterate_retrieve_images()

    dest="./"
    l = [f for f in os.listdir(dest) if f.endswith('.docx')]
    for file in l:
        #convert_single_cell_tables_to_text(file)---Already called in remove_images_from_docx
        final_dest=dest+file
        # related_info(final_dest)
        nested_docx(final_dest)
        stop(final_dest)
    extract_html_md()